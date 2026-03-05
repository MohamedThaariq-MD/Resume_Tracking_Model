import os
import json
from docx import Document
from llm_service import get_llm_client

def get_replacement_mapping(resume_text: str, missing_skills: list[str]) -> dict:
    """
    Uses Gemini to find the optimal words to replace in the resume text with the missing skills.
    Returns a dictionary mapping: {"old_word_to_replace": "new_missing_skill"}
    """
    if not missing_skills:
        return {}
        
        
    
    prompt = f"""
You are an expert resume writer. I have a candidate's resume and a list of skills they are missing for a specific job.
I need to intelligently inject these missing skills into their resume to help them pass the ATS.

Missing Skills to inject: {', '.join(missing_skills)}

Resume Text:
{resume_text}

Analyze the resume and find existing, less important IT skills or generic buzzwords that can be replaced by the missing skills.
The context should still make sense after replacement. 

Return ONLY a valid JSON object where the keys are the EXACT existing words/phrases in the resume (case sensitive, match exactly), and the values are the new missing skills to replace them with.
Limit the replacements to a maximum of {len(missing_skills)} items.
Do not include any markdown formatting, just the raw JSON block.
"""
    try:
        client = get_llm_client()
        if not client:
             return {}
             
        models_to_try = ['gemini-2.0-flash', 'gemini-flash-lite-latest']
        response = None
        
        for index, model_name in enumerate(models_to_try):
            try:
                response = client.models.generate_content(
                    model=model_name,
                    contents=prompt
                )
                break # Success
            except Exception as loop_e:
                if index == len(models_to_try) - 1:
                    raise loop_e # All failed
                else:
                    print(f"Model {model_name} failed. Retrying fallback... ({loop_e})")
                    
        # Parse the JSON response
        text = response.text.replace('```json', '').replace('```', '').strip()
        mapping = json.loads(text)
        return mapping
    except Exception as e:
        print(f"Error generating replacement mapping: {e}")
        return {}

def preserve_format_replace(run, old_text, new_text):
    """ Helper to replace text within a specific run to preserve font/color """
    if old_text in run.text:
        run.text = run.text.replace(old_text, new_text)
        return True
    return False

def magic_edit_docx(input_file_path: str, output_file_path: str, resume_text: str, missing_skills: list[str]) -> bool:
    """
    Edits a .docx file in-place by replacing identified words with missing skills, preserving all formatting.
    Returns True if successful, False otherwise.
    """
    try:
        mapping = get_replacement_mapping(resume_text, missing_skills)
        if not mapping:
            print("No suitable replacements found or missing skills was empty.")
            return False
            
        print(f"Replacement Mapping generated: {mapping}")
            
        doc = Document(input_file_path)
        
        # Iterate over all paragraphs and runs
        for para in doc.paragraphs:
            for old_word, new_skill in mapping.items():
                if old_word in para.text:
                    # To preserve formatting perfectly, we must replace within the 'runs'
                    # Warning: If a word is split across multiple runs, this simple replace might miss it.
                    # For a robust solution, we try to find the run containing the exact word.
                    for run in para.runs:
                        preserve_format_replace(run, old_word, new_skill)
                        
        # Also check tables (often used for formatting resumes)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                         for old_word, new_skill in mapping.items():
                            if old_word in para.text:
                                for run in para.runs:
                                    preserve_format_replace(run, old_word, new_skill)
                                    
        doc.save(output_file_path)
        return True
        
    except Exception as e:
         print(f"Error during magic edit: {e}")
         return False
